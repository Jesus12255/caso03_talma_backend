import base64
from io import BytesIO
import logging
import fitz
import openpyxl
import docx

from utl.file_util import FileUtil

logger = logging.getLogger(__name__)

class FileParser:

    @staticmethod
    def parse(content: bytes, filename: str) -> tuple[list[str], str]:
        if FileUtil.is_valid_pdf(content):
            return FileParser._parse_pdf(content, filename)

        if FileUtil.is_valid_image(content):
            return [base64.b64encode(content).decode()], ""

        if FileUtil.is_valid_xlsx(content):
            return [], FileParser._parse_excel(content, filename)

        if FileUtil.is_valid_docx(content):
            return [], FileParser._parse_word(content, filename)

        return [], ""

    @staticmethod
    def _parse_pdf(content: bytes, filename: str) -> tuple[list[str], str]:
        try:
            images = []
            pdf = fitz.open(stream=BytesIO(content), filetype="pdf")
            for page in pdf:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                images.append(base64.b64encode(pix.tobytes("jpeg")).decode())
            pdf.close()
            return images, ""
        except Exception as e:
            logger.error(f"PDF inválido {filename}: {e}", exc_info=True)
            return [], ""

    @staticmethod
    def _parse_excel(content: bytes, filename: str) -> str:
        try:
            wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"--- HOJA: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(c) for c in row if c is not None]
                    if values:
                        lines.append(",".join(values))
            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Excel inválido {filename}: {e}", exc_info=True)
            return ""

    @staticmethod
    def _parse_word(content: bytes, filename: str) -> str:
        try:
            doc = docx.Document(BytesIO(content))
            lines = []

            for p in doc.paragraphs:
                if p.text.strip():
                    lines.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))

            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Word inválido {filename}: {e}", exc_info=True)
            return ""
